#!/usr/bin/env python3
"""Build MediaKids handbook DOCX files from Markdown sources."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
HANDBOOK_DIR = ROOT / "docs" / "handbooks"
BLUE = RGBColor(0x00, 0x66, 0xCC)
INK = RGBColor(0x1D, 0x1D, 0x1F)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BLUE = "E8F4FC"
FONT = "Noto Sans Thai"


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style_font(style, size: float, color: RGBColor = INK, bold: bool = False) -> None:
    font = style.font
    font.name = FONT
    font.size = Pt(size)
    font.color.rgb = color
    font.bold = bold
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def paragraph_border_bottom(paragraph, color: str = "DADCE0", size: str = "8") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    set_style_font(styles["Normal"], 10.5, INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.18
    set_style_font(styles["Heading 1"], 18, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 2"], 14, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    set_style_font(styles["Heading 3"], 12, INK, True)

    header = section.header.paragraphs[0]
    header.text = "MediaKids Academy Website"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], 8.5, MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Generated for MediaKids website workflow"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], 8.5, MUTED)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("MEDIAKIDS WEBSITE GUIDE")
    set_run_font(run, 9, MUTED, True)
    kicker.paragraph_format.space_after = Pt(4)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, 24, BLUE, True)
    title_p.paragraph_format.space_after = Pt(3)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(subtitle)
    set_run_font(subtitle_run, 11.5, MUTED)
    subtitle_p.paragraph_format.space_after = Pt(14)
    paragraph_border_bottom(subtitle_p, "BBDCF7", "10")

    return doc


def add_inline_markdown(paragraph, text: str, code: bool = False) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLUE
        else:
            set_run_font(run, 10.5, INK)
            if code:
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
                run.font.size = Pt(9)


def add_markdown_content(doc: Document, markdown: str) -> None:
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        shade_paragraph(p, "F5F7FA")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(8)
        add_inline_markdown(p, "\n".join(code_lines), code=True)
        code_lines = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_inline_markdown(p, line[2:])
            continue
        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_inline_markdown(p, number_match.group(1))
            continue
        p = doc.add_paragraph()
        add_inline_markdown(p, line)


def build_docx(markdown_name: str, title: str, subtitle: str, output_name: str) -> None:
    markdown = (HANDBOOK_DIR / markdown_name).read_text(encoding="utf-8")
    doc = setup_document(title, subtitle)
    add_markdown_content(doc, markdown)
    doc.save(HANDBOOK_DIR / output_name)


def main() -> None:
    build_docx(
        "MediaKids-Team-Website-Guide.md",
        "MediaKids Team Website Guide",
        "สำหรับทีมงาน คอมเครื่องอื่น Windows Claude และ AI tools อื่น",
        "MediaKids-Team-Website-Guide.docx",
    )
    build_docx(
        "MediaKids-Mac-Owner-Guide.md",
        "MediaKids Mac Owner Guide",
        "สำหรับเจ้าของเว็บบน Mac เครื่องหลักและ Codex",
        "MediaKids-Mac-Owner-Guide.docx",
    )
    build_docx(
        "MediaKids-Deep-Handoff-Summary.md",
        "MediaKids Deep Handoff Summary",
        "สรุปสิ่งที่เจ้าของเว็บต้องการ สิ่งที่ทำแล้ว และวิธีกู้ระบบถ้าเปลี่ยนเครื่อง",
        "MediaKids-Deep-Handoff-Summary.docx",
    )


if __name__ == "__main__":
    main()
